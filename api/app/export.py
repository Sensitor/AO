"""Export DOCX de la réponse à l'appel d'offres (python-docx)."""
from io import BytesIO

from docx import Document as DocxDocument


def build_response_docx(project_name: str, items: list[dict]) -> bytes:
    """Assemble un .docx. `items` : [{title, content, obligation, verdict}]."""
    doc = DocxDocument()
    doc.add_heading(f"Réponse à l'appel d'offres — {project_name}", level=0)

    if not items:
        doc.add_paragraph("Aucune section générée pour ce projet.")

    for item in items:
        doc.add_heading(item.get("title") or "Exigence", level=1)
        meta_bits = []
        if item.get("obligation"):
            meta_bits.append(f"Niveau : {item['obligation']}")
        if item.get("verdict"):
            meta_bits.append(f"Conformité : {item['verdict']}")
        if meta_bits:
            para = doc.add_paragraph()
            run = para.add_run(" · ".join(meta_bits))
            run.italic = True
        doc.add_paragraph(item.get("content") or "[À compléter]")

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
